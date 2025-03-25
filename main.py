import tkinter as tk
from tkinter import messagebox, ttk
from abc import ABC, abstractmethod
from docx import Document
from openpyxl import Workbook
from pymongo import MongoClient
import threading

# self — это обязательный первый параметр методов класса, который ссылается на текущий экземпляр объекта.

# Абстрактный базовый класс для рецептов
# Иерархия наследования
# Managed-атрибуты meat, cheese, vegetables, sauce
class Recipe(ABC): # Модуль abc позволяет создавать абстрактные классы, которые нельзя инстанцировать напрямую.
    # __init__ - конструктор, задающий обязательные параметры рецепта
    def __init__(self, meat, cheese, vegetables, sauce): # Конструктор класса, инициализирующий обязательные атрибуты рецепта
        # Присваивание переданных значений атрибутам объекта
        # Используется property meat, поэтому значение сохраняется в self._meat
        self.meat = meat
        self.cheese = cheese
        self.vegetables = vegetables
        self.sauce = sauce

    # Декоратор @property для создания геттера атрибута meat. Позволяет только читать атрибут
    @property
    def meat(self):
        # Возвращает значение "защищенного" атрибута _meat
        return self._meat
    # Декоратор @meat.setter определяет метод для установки значения свойства meat. Позволяет изменять атрибут
    @meat.setter
    def meat(self, value):
        # Присваивание переданного значения "защищенному" атрибуту _meat
        # Может включать дополнительную валидацию или преобразование данных
        self._meat = value

    @property
    def cheese(self):
        return self._cheese

    @cheese.setter
    def cheese(self, value):
        self._cheese = value

    @property
    def vegetables(self):
        return self._vegetables

    @vegetables.setter
    def vegetables(self, value):
        self._vegetables = value

    @property
    def sauce(self):
        return self._sauce

    @sauce.setter
    def sauce(self, value):
        self._sauce = value

    @abstractmethod
    def calculate(self):
        pass
    # dunder 1
    # __str__ - определяет строковое представление объекта.
    # Автоматически вызывается функциями str() и print().
    # self - текущий экземпляр класса (передаётся автоматически при вызове)

    # Возвращает:
    #   Строку в формате "ClassName: attr1=value1, attr2=value2,...",
    #   где ClassName - имя класса объекта,
    #   а valueN - значения основных атрибутов объекта.
    # - Предназначен для удобного вывода информации пользователю
    def __str__(self):
        return f"{self.__class__.__name__}: meat={self.meat}, cheese={self.cheese}, vegetables={self.vegetables}, sauce={self.sauce}"
    # 1. self.__class__.__name__ - динамически подставляет имя класса объекта
    #    Пример: "Recipe" для базового класса или "BurgerRecipe" для наследника

    # 2. После двоеточия перечисляются атрибуты в формате "ключ=значение":
    #    meat={self.meat}    - выводит мясной ингредиент
    #    cheese={self.cheese} - выводит сырный ингредиент
    #    vegetables={self.vegetables} - выводит список овощей
    #    sauce={self.sauce}   - выводит соус


    # dunder 2
    # __repr__ - определяет официальное строковое представление объекта.
    # self - текущий экземпляр класса (передаётся автоматически при вызове)

    # Возвращает:
    #   Строку в формате конструктора объекта, содержащую:
    #   - Имя класса (self.__class__.__name__)
    #   - Все значимые атрибуты объекта в виде аргументов
    def __repr__(self):
        return f"{self.__class__.__name__}({self.meat}, {self.cheese}, {self.vegetables}, {self.sauce})"
    # 1. self.__class__.__name__
    #    - Возвращает имя текущего класса как строку (например, "Recipe")
    #
    # 2. {self.meat}, {self.cheese} и т.д.
    #    - Подставляют значения соответствующих атрибутов объекта
    #
    # 3. Общая структура:
    #    - Формирует строку в формате "ClassName(attr1, attr2, attr3, attr4)"
    #
    # Пример вывода для объекта Recipe:
    # "Recipe(говядина, чеддер, ['салат', 'помидор'], кетчуп)"

# Класс для бургера
# Иерархия наследования

# Насследование  ABC (базовый класс Python) -> Recipe (абстрактный класс) -> Burger (конкретная реализация)
class Burger(Recipe):
    def calculate(self):
        # Пример расчета калорий и стоимости для бургера
        # self.meat - инициализированная в конструкторе __init__. Хранит значение, переданное при создании объекта.
        calories = self.meat * 2 + self.cheese * 1.5 + self.vegetables * 0.5 + self.sauce * 0.2
        cost = self.meat * 10 + self.cheese * 5 + self.vegetables * 2 + self.sauce * 1
        return calories, cost

    def __str__(self):
        return f"Burger(meat={self.meat}, cheese={self.cheese}, vegetables={self.vegetables}, sauce={self.sauce})"

    def __repr__(self):
        return f"Burger(meat={self.meat}, cheese={self.cheese}, vegetables={self.vegetables}, sauce={self.sauce})"

# Класс для пиццы
# Иерархия наследования
class Pizza(Recipe):
    def calculate(self):
        # Пример расчета калорий и стоимости для пиццы. Его значение создаётся в GUI интерфейсе.
        calories = self.meat * 1.5 + self.cheese * 2 + self.vegetables * 1 + self.sauce * 0.5
        cost = self.meat * 8 + self.cheese * 6 + self.vegetables * 3 + self.sauce * 2
        return calories, cost

    def __str__(self):
        return f"Pizza(meat={self.meat}, cheese={self.cheese}, vegetables={self.vegetables}, sauce={self.sauce})"

    # def __repr__(self):
    #     return f"Pizza(meat={self.meat}, cheese={self.cheese}, vegetables={self.vegetables}, sauce={self.sauce})"

# Класс для вока
# Иерархия наследования
class Wok(Recipe):
    def calculate(self):
        # Пример расчета калорий и стоимости для вока
        calories = self.meat * 1.8 + self.cheese * 1.2 + self.vegetables * 1.5 + self.sauce * 0.8
        cost = self.meat * 9 + self.cheese * 4 + self.vegetables * 2.5 + self.sauce * 1.5
        return calories, cost

    def __str__(self):
        return f"Wok(meat={self.meat}, cheese={self.cheese}, vegetables={self.vegetables}, sauce={self.sauce})"

    def __repr__(self):
        return f"Wok(meat={self.meat}, cheese={self.cheese}, vegetables={self.vegetables}, sauce={self.sauce})"

# Функция для сохранения отчета в формате .docx
def save_docx(recipe_name, calories, cost, filename):
    doc = Document()
    doc.add_heading(f'Отчет по рецепту: {recipe_name}', 0)
    doc.add_paragraph(f'Энергетическая ценность: {calories} ккал')
    doc.add_paragraph(f'Стоимость: {cost} руб.')
    doc.save(filename)

# Функция для сохранения отчета в формате .xlsx
def save_xlsx(recipe_name, calories, cost, filename):
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет"
    ws['A1'] = 'Рецепт'
    ws['B1'] = 'Калории (ккал)'
    ws['C1'] = 'Стоимость (руб.)'
    ws['A2'] = recipe_name
    ws['B2'] = calories
    ws['C2'] = cost
    wb.save(filename)

# Функция для сохранения в MongoDB
def save_to_mongodb(recipe_name, calories, cost):
    client = MongoClient("mongodb://localhost:27017/")
    db = client["recipes_db"]
    collection = db["recipes"]
    collection.insert_one({
        "name": recipe_name,
        "calories": calories,
        "cost": cost
    })
    client.close()

# Функция для обработки ввода и расчета
def on_calculate():
    def calculate():
        recipe_name = recipe_var.get()
        ingredients = {
            "meat": float(meat_entry.get()),
            "cheese": float(cheese_entry.get()),
            "vegetables": float(vegetables_entry.get()),
            "sauce": float(sauce_entry.get())
        }

        if recipe_name == "Бургер":
            recipe = Burger(**ingredients)
        elif recipe_name == "Пицца":
            recipe = Pizza(**ingredients)
            print(repr(recipe))
        elif recipe_name == "Вок":
            recipe = Wok(**ingredients)
            print(str(recipe))
            print(repr(recipe))


        calories, cost = recipe.calculate()

        calories_label.config(text=f"Энергетическая ценность: {calories} ккал")
        cost_label.config(text=f"Стоимость: {cost} руб.")

        report_format = report_var.get()
        if report_format == "Word (.docx)":
            filename = f"{recipe_name}_report.docx"
            save_docx(recipe_name, calories, cost, filename)
        elif report_format == "Excel (.xlsx)":
            filename = f"{recipe_name}_report.xlsx"
            save_xlsx(recipe_name, calories, cost, filename)

        save_to_mongodb(recipe_name, calories, cost)
        messagebox.showinfo("Сохранено", f"Отчет сохранен в файл {filename} и в MongoDB")

    threading.Thread(target=calculate).start()

# Создание графического интерфейса с использованием Tkinter
root = tk.Tk()
root.title("Рецепты")
root.geometry("400x400")

recipe_var = tk.StringVar()
report_var = tk.StringVar()

tk.Label(root, text="Рецепт").grid(row=0, column=0)
recipe_combobox = ttk.Combobox(root, textvariable=recipe_var, values=["Бургер", "Пицца", "Вок"])
recipe_combobox.grid(row=0, column=1)

tk.Label(root, text="Мясо").grid(row=1, column=0)
meat_entry = tk.Entry(root)
meat_entry.grid(row=1, column=1)

tk.Label(root, text="Сыр").grid(row=2, column=0)
cheese_entry = tk.Entry(root)
cheese_entry.grid(row=2, column=1)

tk.Label(root, text="Овощи").grid(row=3, column=0)
vegetables_entry = tk.Entry(root)
vegetables_entry.grid(row=3, column=1)

tk.Label(root, text="Соус").grid(row=4, column=0)
sauce_entry = tk.Entry(root)
sauce_entry.grid(row=4, column=1)

tk.Label(root, text="Формат отчета").grid(row=5, column=0)
report_combobox = ttk.Combobox(root, textvariable=report_var, values=["Word (.docx)", "Excel (.xlsx)"])
report_combobox.grid(row=5, column=1)

calculate_button = tk.Button(root, text="Рассчитать", command=on_calculate)
calculate_button.grid(row=6, column=0, columnspan=2)

calories_label = tk.Label(root, text="Энергетическая ценность: ")
calories_label.grid(row=7, column=0, columnspan=2)

cost_label = tk.Label(root, text="Стоимость: ")
cost_label.grid(row=8, column=0, columnspan=2)

root.mainloop()